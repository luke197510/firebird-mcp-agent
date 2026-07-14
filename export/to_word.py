import io
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def artifact_to_word(text: str, dataframes: list[pd.DataFrame]) -> bytes:
    """Genera un file .docx con il testo e le tabelle dell'artifact."""
    doc = Document()

    for line in text.splitlines():
        if re.match(r"^### ", line):
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^## ", line):
            doc.add_heading(line[3:].strip(), level=2)
        elif re.match(r"^# ", line):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip():
            doc.add_paragraph(line.strip("* \t"))

    for i, df in enumerate(dataframes):
        doc.add_heading(f"Tabella {i + 1}", level=3)
        if df.empty:
            continue
        table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
        table.style = "Table Grid"

        # Header
        hdr = table.rows[0]
        for j, col in enumerate(df.columns):
            cell = hdr.cells[j]
            cell.text = str(col)
            run = cell.paragraphs[0].runs[0]
            run.bold = True

        # Dati
        for r_i, row in df.iterrows():
            for c_i, val in enumerate(row):
                table.rows[r_i + 1].cells[c_i].text = str(val)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
